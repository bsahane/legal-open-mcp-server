"""Document ingest, search and review tools for the Legal MCP Server.

Ingested documents live in the user's own PostgreSQL instance. Chunking follows
the document's own structure - numbered clauses, section headings, paragraph
breaks - rather than a fixed window, because a contract clause split down the
middle retrieves badly and reads worse.

Search is hybrid: Postgres full-text and pgvector cosine similarity, fused by
reciprocal rank. When embeddings are unavailable the tool says so in its
response instead of quietly returning weaker results.
"""

import hashlib
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from legal_mcp_server.src.domain import clause_rules
from legal_mcp_server.src.sources import embeddings as emb
from legal_mcp_server.src.storage.legal_store import (
    get_store,
    unavailable_response,
)
from legal_mcp_server.utils.pylogger import get_python_logger

logger = get_python_logger()

MAX_CHUNK_CHARS = 1800
MIN_CHUNK_CHARS = 200
RRF_K = 60

SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}

# Clause and section starts: "1.", "1.1", "(a)", "Clause 5", "Section 12", "ARTICLE IV".
_STRUCTURE_BREAK = re.compile(
    r"^\s*(?:"
    r"\d{1,2}(?:\.\d{1,2}){0,3}[.)]?\s+[A-Z(]"
    r"|\([a-z]{1,3}\)\s+[A-Z]"
    r"|(?:CLAUSE|Clause|SECTION|Section|ARTICLE|Article)\s+[\dIVXLC]+"
    r"|[A-Z][A-Z \-]{6,60}$"
    r")",
    re.MULTILINE,
)


def extract_text(path: Path) -> Dict[str, Any]:
    """Extract text from a PDF, DOCX, TXT or Markdown file.

    Args:
        path: Path to the file.

    Returns:
        Dict with ``text`` and ``page_count``.

    Raises:
        ValueError: If the format is unsupported or the file yields no text.
    """
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages)
        if not text.strip():
            raise ValueError(
                f"No text could be extracted from {path.name}. It is probably a "
                "scanned image. Run OCR over it first - this server does not "
                "perform OCR, and an empty document indexed silently is worse "
                "than a failed ingest."
            )
        return {"text": text, "page_count": len(pages)}

    if suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        paragraphs = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                paragraphs.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(paragraphs)
        if not text.strip():
            raise ValueError(f"No text found in {path.name}.")
        return {"text": text, "page_count": None}

    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        if not text.strip():
            raise ValueError(f"{path.name} is empty.")
        return {"text": text, "page_count": None}

    raise ValueError(
        f"Unsupported file type '{suffix}'. Supported: "
        f"{', '.join(sorted(SUPPORTED_SUFFIXES))}."
    )


def chunk_text(text: str) -> List[Dict[str, Any]]:
    """Split a document at its own structural boundaries.

    Args:
        text: The full document text.

    Returns:
        List of chunks with ``content`` and a ``heading_path`` hint.
    """
    breaks = [m.start() for m in _STRUCTURE_BREAK.finditer(text)]
    boundaries = sorted({0, *breaks, len(text)})

    segments: List[str] = []
    for start, end in zip(boundaries, boundaries[1:]):
        segment = text[start:end].strip()
        if segment:
            segments.append(segment)

    if not segments:
        segments = [text.strip()]

    # Merge segments that are too small, split those that are too large.
    chunks: List[str] = []
    buffer = ""
    for segment in segments:
        if len(segment) > MAX_CHUNK_CHARS:
            if buffer:
                chunks.append(buffer)
                buffer = ""
            paragraphs = re.split(r"\n\s*\n", segment)
            current = ""
            for paragraph in paragraphs:
                if len(current) + len(paragraph) > MAX_CHUNK_CHARS and current:
                    chunks.append(current.strip())
                    current = paragraph
                else:
                    current = f"{current}\n\n{paragraph}" if current else paragraph
            if current.strip():
                chunks.append(current.strip())
        elif len(buffer) + len(segment) < MIN_CHUNK_CHARS:
            buffer = f"{buffer}\n\n{segment}" if buffer else segment
        else:
            if buffer:
                chunks.append(buffer)
            buffer = segment

    if buffer:
        chunks.append(buffer)

    result = []
    for index, content in enumerate(chunks):
        first_line = content.split("\n", 1)[0].strip()
        result.append(
            {
                "chunk_index": index,
                "content": content,
                "heading_path": first_line[:180] if len(first_line) < 180 else None,
            }
        )
    return result


async def ingest_document(
    file_path: str,
    title: Optional[str] = None,
    doc_type: Optional[str] = None,
    matter_id: Optional[int] = None,
) -> Dict[str, Any]:
    """Ingest a document into the searchable store.

    TOOL_NAME=ingest_document
    DISPLAY_NAME=Ingest Document
    USECASE=Add a contract, notice, judgment, pleading or case bundle to the searchable index so it can be questioned later
    INSTRUCTIONS=1. Give the absolute path to a PDF, DOCX, TXT or MD file, 2. Link it to a matter where one exists, 3. If ingest reports a scanned PDF, OCR it before retrying - it will not be indexed blind
    INPUT_DESCRIPTION=file_path (string, required): absolute path to the file. title (string, optional): defaults to the filename. doc_type (string, optional): e.g. "contract", "judgment", "notice", "pleading". matter_id (int, optional): matter to attach it to.
    OUTPUT_DESCRIPTION=Dictionary with status, the document id, chunk count, page count, whether embeddings were generated, and a duplicate flag if the file was already ingested
    EXAMPLES=ingest_document("/Users/me/docs/vendor-agreement.pdf", doc_type="contract", matter_id=3)
    PREREQUISITES=PostgreSQL running. Semantic search additionally needs EMBEDDING_PROVIDER set to voyage or local.
    RELATED_TOOLS=search_my_documents to query it; review_contract to assess it; extract_clauses to map it

    I/O-bound operation - uses async def for file and database access.

    Args:
        file_path: Absolute path to the source file.
        title: Display title; defaults to the filename.
        doc_type: Category of document.
        matter_id: Matter to attach the document to.

    Returns:
        Dict describing the ingested document.
    """
    try:
        path = Path(file_path).expanduser()
        if not path.is_file():
            raise ValueError(f"No file at {path}")
        if path.suffix.lower() not in SUPPORTED_SUFFIXES:
            raise ValueError(
                f"Unsupported file type '{path.suffix}'. Supported: "
                f"{', '.join(sorted(SUPPORTED_SUFFIXES))}."
            )

        raw = path.read_bytes()
        digest = hashlib.sha256(raw).hexdigest()

        store = get_store()
        existing = await store.fetchrow(
            "SELECT id, title FROM documents WHERE sha256 = $1", digest
        )
        if existing is not None:
            return {
                "status": "success",
                "operation": "ingest_document",
                "document_id": existing["id"],
                "title": existing["title"],
                "duplicate": True,
                "message": (
                    f"This file is already ingested as document "
                    f"#{existing['id']} ('{existing['title']}'). Nothing was "
                    "re-indexed."
                ),
            }

        extracted = extract_text(path)
        text = extracted["text"]
        chunks = chunk_text(text)

        document_id = await store.fetchval(
            """
            INSERT INTO documents (matter_id, title, doc_type, source_path, sha256,
                                   page_count, char_count, full_text)
            VALUES ($1,$2,$3,$4,$5,$6,$7,$8) RETURNING id
            """,
            matter_id,
            title or path.stem,
            doc_type,
            str(path),
            digest,
            extracted["page_count"],
            len(text),
            text,
        )

        vectors: Optional[List[List[float]]] = None
        embedding_note: Optional[str] = None
        try:
            provider = emb.get_provider()
            if provider is not None:
                vectors = await provider.embed_documents([c["content"] for c in chunks])
        except emb.EmbeddingUnavailable as e:
            embedding_note = (
                f"Embeddings were not generated ({e}). The document is indexed for "
                "full-text search only; semantic queries over it will be weak."
            )
            logger.warning(f"Embedding failed during ingest: {e}")

        for index, chunk in enumerate(chunks):
            if vectors is not None and store.vector_enabled:
                await store.execute(
                    "INSERT INTO document_chunks (document_id, chunk_index, "
                    "heading_path, content, embedding) VALUES ($1,$2,$3,$4,$5)",
                    document_id,
                    chunk["chunk_index"],
                    chunk["heading_path"],
                    chunk["content"],
                    str(vectors[index]),
                )
            else:
                await store.execute(
                    "INSERT INTO document_chunks (document_id, chunk_index, "
                    "heading_path, content) VALUES ($1,$2,$3,$4)",
                    document_id,
                    chunk["chunk_index"],
                    chunk["heading_path"],
                    chunk["content"],
                )

        logger.info(f"Ingested document {document_id} with {len(chunks)} chunks")

        return {
            "status": "success",
            "operation": "ingest_document",
            "document_id": document_id,
            "title": title or path.stem,
            "doc_type": doc_type,
            "matter_id": matter_id,
            "page_count": extracted["page_count"],
            "char_count": len(text),
            "chunk_count": len(chunks),
            "embeddings_generated": vectors is not None and store.vector_enabled,
            "embedding_note": embedding_note,
            "duplicate": False,
            "message": (
                f"Ingested '{title or path.stem}' as document #{document_id} "
                f"in {len(chunks)} chunks."
                + (f" {embedding_note}" if embedding_note else "")
            ),
        }

    except ConnectionError as e:
        return unavailable_response("ingest_document", e)
    except Exception as e:
        logger.error(f"Error in ingest_document: {e}")
        return {
            "status": "error",
            "operation": "ingest_document",
            "error": str(e),
            "message": f"Failed to ingest document: {e}",
        }


async def search_my_documents(
    query: str,
    matter_id: Optional[int] = None,
    doc_type: Optional[str] = None,
    limit: int = 8,
) -> Dict[str, Any]:
    """Search your own ingested documents.

    TOOL_NAME=search_my_documents
    DISPLAY_NAME=Search My Documents
    USECASE=Find the passage in your own contracts, notices, judgments or case bundles that answers a question
    INSTRUCTIONS=1. Ask for the substance you need, not just keywords, 2. Filter by matter where you know it, 3. Check search_mode in the response - 'fulltext_only' means conceptual queries were not really answered
    INPUT_DESCRIPTION=query (string, required): what you are looking for. matter_id (int, optional): restrict to one matter. doc_type (string, optional): restrict by category. limit (int, optional, default 8): maximum passages.
    OUTPUT_DESCRIPTION=Dictionary with status, matching passages with document title, heading and content, the search_mode actually used, and a note when semantic search was unavailable
    EXAMPLES=search_my_documents("termination for convenience notice period"), search_my_documents("indemnity cap", matter_id=3)
    PREREQUISITES=PostgreSQL running with documents ingested. Semantic ranking needs an embedding provider.
    RELATED_TOOLS=ingest_document to add documents; get_document to read one in full

    I/O-bound operation - uses async def for database access.

    Args:
        query: What to search for.
        matter_id: Optional matter filter.
        doc_type: Optional document-type filter.
        limit: Maximum passages to return.

    Returns:
        Dict with the matching passages.
    """
    try:
        if not query or not query.strip():
            raise ValueError("query must be a non-empty string")

        bounded = max(1, min(limit, 50))
        store = get_store()

        filters = []
        params: List[Any] = []
        if matter_id is not None:
            params.append(matter_id)
            filters.append(f"d.matter_id = ${len(params)}")
        if doc_type:
            params.append(doc_type)
            filters.append(f"d.doc_type = ${len(params)}")
        where = f"AND {' AND '.join(filters)}" if filters else ""

        # Full-text leg.
        params_ft = [*params, query.strip(), bounded * 3]
        text_rows = await store.fetch(
            f"""
            SELECT c.id, c.document_id, c.chunk_index, c.heading_path, c.content,
                   d.title, d.doc_type, d.matter_id,
                   ts_rank(c.tsv, plainto_tsquery('english', ${len(params) + 1})) AS score
            FROM document_chunks c
            JOIN documents d ON d.id = c.document_id
            WHERE c.tsv @@ plainto_tsquery('english', ${len(params) + 1}) {where}
            ORDER BY score DESC
            LIMIT ${len(params) + 2}
            """,
            *params_ft,
        )

        # Vector leg.
        vector_rows: List[Any] = []
        search_mode = "fulltext_only"
        semantic_note: Optional[str] = None

        try:
            provider = emb.get_provider()
            if provider is not None and store.vector_enabled:
                vector = await provider.embed_query(query.strip())
                params_vec = [*params, str(vector), bounded * 3]
                vector_rows = await store.fetch(
                    f"""
                    SELECT c.id, c.document_id, c.chunk_index, c.heading_path,
                           c.content, d.title, d.doc_type, d.matter_id,
                           1 - (c.embedding <=> ${len(params) + 1}::vector) AS score
                    FROM document_chunks c
                    JOIN documents d ON d.id = c.document_id
                    WHERE c.embedding IS NOT NULL {where}
                    ORDER BY c.embedding <=> ${len(params) + 1}::vector
                    LIMIT ${len(params) + 2}
                    """,
                    *params_vec,
                )
                search_mode = "hybrid"
            elif provider is None:
                semantic_note = emb.provider_status()["note"]
            else:
                semantic_note = (
                    "pgvector is not enabled on this database, so only full-text "
                    "matching was used."
                )
        except emb.EmbeddingUnavailable as e:
            semantic_note = f"Semantic search unavailable ({e}); full-text only."
            logger.warning(f"Vector search skipped: {e}")

        # Reciprocal rank fusion.
        fused: Dict[int, Dict[str, Any]] = {}
        for rank, row in enumerate(text_rows):
            entry = fused.setdefault(
                row["id"], {"row": row, "rrf": 0.0, "matched_by": []}
            )
            entry["rrf"] += 1.0 / (RRF_K + rank + 1)
            entry["matched_by"].append("fulltext")
        for rank, row in enumerate(vector_rows):
            entry = fused.setdefault(
                row["id"], {"row": row, "rrf": 0.0, "matched_by": []}
            )
            entry["rrf"] += 1.0 / (RRF_K + rank + 1)
            entry["matched_by"].append("semantic")

        ordered = sorted(fused.values(), key=lambda e: -e["rrf"])[:bounded]

        results = []
        for entry in ordered:
            row = entry["row"]
            results.append(
                {
                    "chunk_id": row["id"],
                    "document_id": row["document_id"],
                    "document_title": row["title"],
                    "doc_type": row["doc_type"],
                    "matter_id": row["matter_id"],
                    "chunk_index": row["chunk_index"],
                    "heading": row["heading_path"],
                    "content": row["content"],
                    "relevance": round(entry["rrf"], 5),
                    "matched_by": entry["matched_by"],
                }
            )

        return {
            "status": "success",
            "operation": "search_my_documents",
            "query": query,
            "results": results,
            "result_count": len(results),
            "search_mode": search_mode,
            "semantic_note": semantic_note,
            "message": (
                f"{len(results)} passages found using {search_mode} search."
                + (f" {semantic_note}" if semantic_note else "")
                + (
                    ""
                    if results
                    else " Nothing matched. If the documents are scanned PDFs they "
                    "may have been ingested without text."
                )
            ),
        }

    except ConnectionError as e:
        return unavailable_response("search_my_documents", e)
    except Exception as e:
        logger.error(f"Error in search_my_documents: {e}")
        return {
            "status": "error",
            "operation": "search_my_documents",
            "error": str(e),
            "message": "Failed to search documents",
        }


async def list_my_documents(
    matter_id: Optional[int] = None, limit: int = 50
) -> Dict[str, Any]:
    """List ingested documents.

    TOOL_NAME=list_my_documents
    DISPLAY_NAME=List My Documents
    USECASE=See what has been ingested before searching, so an empty search result can be told apart from an empty index
    INSTRUCTIONS=1. Call with no filter for everything, or a matter_id to scope it
    INPUT_DESCRIPTION=matter_id (int, optional). limit (int, optional, default 50).
    OUTPUT_DESCRIPTION=Dictionary with status, documents with id, title, type, page and chunk counts, and the total
    EXAMPLES=list_my_documents(), list_my_documents(matter_id=3)
    PREREQUISITES=PostgreSQL running
    RELATED_TOOLS=ingest_document, search_my_documents, get_document

    I/O-bound operation - uses async def for database access.

    Args:
        matter_id: Optional matter filter.
        limit: Maximum documents to return.

    Returns:
        Dict with the document list.
    """
    try:
        store = get_store()
        bounded = max(1, min(limit, 500))

        if matter_id is not None:
            rows = await store.fetch(
                """
                SELECT d.id, d.title, d.doc_type, d.matter_id, d.page_count,
                       d.char_count, d.created_at,
                       (SELECT count(*) FROM document_chunks c WHERE c.document_id = d.id) AS chunks
                FROM documents d WHERE d.matter_id = $1
                ORDER BY d.created_at DESC LIMIT $2
                """,
                matter_id,
                bounded,
            )
        else:
            rows = await store.fetch(
                """
                SELECT d.id, d.title, d.doc_type, d.matter_id, d.page_count,
                       d.char_count, d.created_at,
                       (SELECT count(*) FROM document_chunks c WHERE c.document_id = d.id) AS chunks
                FROM documents d ORDER BY d.created_at DESC LIMIT $1
                """,
                bounded,
            )

        documents = [
            {**dict(r), "created_at": r["created_at"].isoformat()} for r in rows
        ]

        return {
            "status": "success",
            "operation": "list_my_documents",
            "documents": documents,
            "document_count": len(documents),
            "embedding_status": emb.provider_status(),
            "message": f"{len(documents)} documents ingested.",
        }

    except ConnectionError as e:
        return unavailable_response("list_my_documents", e)
    except Exception as e:
        logger.error(f"Error in list_my_documents: {e}")
        return {
            "status": "error",
            "operation": "list_my_documents",
            "error": str(e),
            "message": "Failed to list documents",
        }


async def get_document(document_id: int, max_chars: int = 20000) -> Dict[str, Any]:
    """Retrieve an ingested document's text.

    TOOL_NAME=get_document
    DISPLAY_NAME=Get Document
    USECASE=Read a whole ingested document rather than isolated passages
    INSTRUCTIONS=1. Give the document id from list_my_documents or a search result, 2. Raise max_chars if the document is truncated and you need the rest
    INPUT_DESCRIPTION=document_id (int, required). max_chars (int, optional, default 20000): truncation limit; 0 for the whole document.
    OUTPUT_DESCRIPTION=Dictionary with status, the document metadata, its text, and whether the text was truncated
    EXAMPLES=get_document(5), get_document(5, max_chars=0)
    PREREQUISITES=PostgreSQL running
    RELATED_TOOLS=search_my_documents to find the relevant part of a long document instead

    I/O-bound operation - uses async def for database access.

    Args:
        document_id: The document to retrieve.
        max_chars: Truncation limit; 0 returns everything.

    Returns:
        Dict with the document text.
    """
    try:
        if not isinstance(document_id, int) or document_id <= 0:
            raise ValueError("document_id must be a positive integer")

        store = get_store()
        row = await store.fetchrow("SELECT * FROM documents WHERE id = $1", document_id)
        if row is None:
            return {
                "status": "not_found",
                "operation": "get_document",
                "document_id": document_id,
                "message": f"No document with id {document_id}.",
            }

        text = row["full_text"] or ""
        truncated = bool(max_chars) and len(text) > max_chars
        if truncated:
            text = text[:max_chars]

        return {
            "status": "success",
            "operation": "get_document",
            "document_id": row["id"],
            "title": row["title"],
            "doc_type": row["doc_type"],
            "matter_id": row["matter_id"],
            "page_count": row["page_count"],
            "char_count": row["char_count"],
            "source_path": row["source_path"],
            "text": text,
            "truncated": truncated,
            "message": (
                f"Document #{document_id}: {row['title']}."
                + (
                    f" Text truncated at {max_chars} of {row['char_count']} "
                    "characters; use search_my_documents to find a specific "
                    "passage, or raise max_chars."
                    if truncated
                    else ""
                )
            ),
        }

    except ConnectionError as e:
        return unavailable_response("get_document", e)
    except Exception as e:
        logger.error(f"Error in get_document: {e}")
        return {
            "status": "error",
            "operation": "get_document",
            "error": str(e),
            "message": "Failed to retrieve document",
        }


def extract_clauses(text: str) -> Dict[str, Any]:
    """Map a contract's clauses onto a standard taxonomy.

    TOOL_NAME=extract_clauses
    DISPLAY_NAME=Extract Contract Clauses
    USECASE=See at a glance what a contract does and does not provide for, before reading it line by line
    INSTRUCTIONS=1. Pass the contract text, 2. Read missing_expected_clauses first - what a contract omits is usually more dangerous than what it says
    INPUT_DESCRIPTION=text (string, required): the contract text. Use get_document to obtain it for an ingested file.
    OUTPUT_DESCRIPTION=Dictionary with status, clauses found by category with excerpts and positions, the categories a commercial contract normally has but this one lacks, and counts
    EXAMPLES=extract_clauses(contract_text)
    PREREQUISITES=None - fully offline
    RELATED_TOOLS=review_contract for the risk assessment; get_document to fetch the text

    CPU-bound operation - uses def for pattern matching.

    Args:
        text: The contract text.

    Returns:
        Dict with the clause map.
    """
    try:
        if not text or not text.strip():
            raise ValueError("text must be a non-empty string")

        found = clause_rules.classify_clauses(text)
        labels = {c.key: c.label for c in clause_rules.CLAUSE_TAXONOMY}

        clauses = [
            {
                "type": key,
                "label": labels.get(key, key),
                "occurrences": len(hits),
                "excerpts": [excerpt for _, excerpt in hits[:2]],
                "first_position": hits[0][0],
            }
            for key, hits in sorted(found.items(), key=lambda kv: kv[1][0][0])
        ]

        missing = clause_rules.missing_expected_clauses(text)

        return {
            "status": "success",
            "operation": "extract_clauses",
            "clauses": clauses,
            "clause_count": len(clauses),
            "missing_expected_clauses": [
                {"type": c.key, "label": c.label} for c in missing
            ],
            "message": (
                f"Identified {len(clauses)} clause categories."
                + (
                    f" {len(missing)} categories a commercial contract would "
                    f"normally contain are absent: "
                    f"{', '.join(c.label for c in missing)}."
                    if missing
                    else ""
                )
                + " Categories are matched by pattern, so read the excerpts rather "
                "than trusting the labels blindly."
            ),
        }

    except Exception as e:
        logger.error(f"Error in extract_clauses: {e}")
        return {
            "status": "error",
            "operation": "extract_clauses",
            "error": str(e),
            "message": "Failed to extract clauses",
        }


def review_contract(text: str, party_side: Optional[str] = None) -> Dict[str, Any]:
    """Flag risky and missing provisions in a contract under Indian law.

    TOOL_NAME=review_contract
    DISPLAY_NAME=Contract Risk Review
    USECASE=Get a structured first-pass risk review of a contract, grounded in Indian statutory law rather than generic commercial advice
    INSTRUCTIONS=1. Pass the contract text, 2. Work through the high-severity flags first, 3. Treat each flag as a prompt to read the clause, not as a conclusion, 4. Note that section 27 makes most post-termination non-competes void in India regardless of how reasonable they look
    INPUT_DESCRIPTION=text (string, required): the contract text. party_side (string, optional): which side you act for, to frame the summary.
    OUTPUT_DESCRIPTION=Dictionary with status, risk flags by severity with the statutory authority for each, the clause map, missing expected clauses, and a summary
    EXAMPLES=review_contract(contract_text), review_contract(contract_text, party_side="the employee")
    PREREQUISITES=None - fully offline
    RELATED_TOOLS=extract_clauses for the clause map alone; get_section to read any provision cited; search_case_law for judicial treatment

    CPU-bound operation - uses def for rule evaluation.

    Args:
        text: The contract text.
        party_side: Which party's perspective to frame the review from.

    Returns:
        Dict with the risk review.
    """
    try:
        if not text or not text.strip():
            raise ValueError("text must be a non-empty string")

        flags = clause_rules.assess_risks(text)
        by_severity: Dict[str, List[Dict[str, Any]]] = {}
        for flag in flags:
            by_severity.setdefault(flag.severity.value, []).append(flag.to_dict())

        found = clause_rules.classify_clauses(text)
        missing = clause_rules.missing_expected_clauses(text)
        high = len(by_severity.get("high", []))

        return {
            "status": "success",
            "operation": "review_contract",
            "party_side": party_side,
            "flags": [f.to_dict() for f in flags],
            "flags_by_severity": by_severity,
            "flag_count": len(flags),
            "high_severity_count": high,
            "clause_types_present": sorted(found),
            "missing_expected_clauses": [
                {"type": c.key, "label": c.label} for c in missing
            ],
            "message": (
                f"{len(flags)} issues flagged, {high} of them high severity."
                + (f" Reviewing on behalf of {party_side}." if party_side else "")
                + " These are pattern-based checks that tell you which clauses to "
                "read; they are not an opinion on enforceability. Confirm anything "
                "material against the provision cited and, where it matters, "
                "against case law."
            ),
            "disclaimer": (
                "A rule-based review cannot catch everything and cannot weigh "
                "commercial context. It does not replace review by an advocate "
                "for a contract of real value."
            ),
        }

    except Exception as e:
        logger.error(f"Error in review_contract: {e}")
        return {
            "status": "error",
            "operation": "review_contract",
            "error": str(e),
            "message": "Failed to review contract",
        }


TOOLS: List[Any] = [
    ingest_document,
    search_my_documents,
    list_my_documents,
    get_document,
    extract_clauses,
    review_contract,
]
